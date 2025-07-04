# DONE: Add option to provide https link to extract susbtitles from.
# TODO: Handle input out of range
# DONE: if we are passed a video file name just process it.......

import nltk
import os
import re
import json
import argparse
import urllib.request
# import textwrap
import ffmpeg
import subprocess
import sys
import requests
import html2markdown
from docx import Document
from argparse import RawTextHelpFormatter
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module='html2markdown')

url = "https://b.jw-cdn.org/apis/mediator/v1/categories/E/LatestVideos?detailed=1&clientType=www"
tmp = ""
banner = """

     ██╗██╗    ██╗    ███████╗██╗   ██╗██████╗ ████████╗██╗████████╗██╗     ███████╗███████╗
     ██║██║    ██║    ██╔════╝██║   ██║██╔══██╗╚══██╔══╝██║╚══██╔══╝██║     ██╔════╝██╔════╝
     ██║██║ █╗ ██║    ███████╗██║   ██║██████╔╝   ██║   ██║   ██║   ██║     █████╗  ███████╗
██   ██║██║███╗██║    ╚════██║██║   ██║██╔══██╗   ██║   ██║   ██║   ██║     ██╔══╝  ╚════██║
╚█████╔╝╚███╔███╔╝    ███████║╚██████╔╝██████╔╝   ██║   ██║   ██║   ███████╗███████╗███████║
 ╚════╝  ╚══╝╚══╝     ╚══════╝ ╚═════╝ ╚═════╝    ╚═╝   ╚═╝   ╚═╝   ╚══════╝╚══════╝╚══════╝
                                                                                            
"""

# """
#   .-..-.  .-.     .----..-. .-.----. .-----..-..-----..-.   .----. .----.
#   | || {  } |    { {__-`| } { | {_} }`-' '-'{ |`-' '-'} |   } |__}{ {__-`
# {`-' }{  /\  }    .-._} }\ `-' / {_} }  } {  | }  } {  } '--.} '__}.-._} }
# `---'`-'  `-'    `----'  `---'`----'   `-'  `-'  `-'  `----'`----'`----'
#                                                                           """
helptext = "help text"
# """
# This program has two modes of operation:

#             With no command line parameters the program queries for the latest videos from jw.org and allows
#             you to choose which one you'd like the subtitles for and then displays them.

#             If you provide the URL of a video file as the command line parameter, the program downloads the video
#             and attempts to extract the subtitle information from the video to display.

#             There are two additional parameters that can passed when providing the URL:
#                 -t Outputs the subtitles to a text file in the current directory named with the
#                     title of the video.
#                 -d Outputs the subtitles to a Microsoft Word docx file in the current directory
#                     named with the title of the video.
#             IMPORTANT: The second mode requires that you have installed FFpmeg and the ffmpeg.exe program is
#             available in your PATH.

#             You can find FFmpeg at:  https://ffmpeg.org/download.html
# """
helptextalt = """               This program queries for the latest videos from jw.org and allows
            you to choose which one you'd like the subtitles for and then displays them.
            """


def help():
    print(banner)
    print(main.__doc__)
    sys.exit()


def download_file(url):
    """
    download_file accepts a url passed, uses the requests library to obtain the document and returns the results as lines.
    Designed to work with html documents primarily."""
    r = requests.get(url)
    lines = r.iter_lines()
    return lines


def download_video(url, mytmp):
    filename = mytmp + url.split("/")[-1]
    print("Filename: ", filename)
    if os.path.isfile(filename):
        print("Already downloaded: ", filename)
        return (filename)
    print("Downloading to: ", filename)
    print("Checking for video file at URL provided.")
    with urllib.request.urlopen(url) as response:
        info = response.info()
        print("File type found is: ", info.get_content_maintype())
    if info.get_content_maintype() != "video":
        print("Video file not found at ", url)
        return ("fail")
    try:
        query_parameters = {"downloadformat": "mp4"}
        response = requests.get(url, params=query_parameters, stream=True)
        if response.ok:
            with open(filename, mode="wb") as file:
                for chunk in response.iter_content(chunk_size=10 * 1024):
                    file.write(chunk)
        # else:
        #     print("Something didn't work downloading the file.")
        #     print("Error: ", response.ok)
        #     return("fail")
        return (filename)
    except Exception as e:
        print(e)
        return ("fail")


def get_title(input_file):
    """
    The get_title function accepts the name of a video file containing an embedded title in its metdata, extracts
    it with ffmpeg and returns it.
    """
    try:
        metadata = ffmpeg.probe(input_file, show_entries="format_tags=title")
        title = metadata.get('format', {}).get('tags', {}).get('title', '')
        title = re.sub(":", "\:", title)
        title = re.sub('"', "", title)
        # title = re.sub("\.", "", title)
        # title = re.sub("(", "- ", title)
        # title = re.sub(")", "", "title)
        title = re.sub(",", "", title)
        if title:
            return title.strip()
        else:
            return ()
    except ffmpeg.Error as e:
        print(f"Error getting title metadata: {e.stderr}")
        return None


def break_into_paragraphs_nltk(text):
    # strip out ', '
    text = re.sub("', '", "", text)
    # Use nltk's sent_tokenize and then join into paragraphs
    sentences = nltk.sent_tokenize(text)
    paragraphs = []
    current_paragraph = []

    for sentence in sentences:
        if sentence == '':
            continue
        current_paragraph.append(sentence)
        # Create a new paragraph after 5 sentences
        if len(current_paragraph) >= 5:
            paragraphs.append(' '.join(current_paragraph))
            current_paragraph = []

    if current_paragraph:  # add remaining sentences
        paragraphs.append(' '.join(current_paragraph))
    for i, paragraph in enumerate(paragraphs):
        print(f"{paragraph}\n  ")
    return paragraphs


def get_subtitles(video):
    """
    The get_subtitles function accepts the name of a video file with embedded subtitle metadata, extracts
    the subtitle metadata with ffmpeg and returns the subtitles.
    """
    srt = video + ".srt"
    # print("Processing subtitles from video to:", srt)
    cmd = ["ffmpeg", "-y", "-i", video, srt]
    # os.remove(srt)
    try:
        subprocess.run(cmd, check=True, capture_output=True)
    except CalledProcessError as e:
        print(f"Error extracting subtitles: {e.stderr}")
    with open(srt, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    return (lines)


def processlines(lines):
    """
    The processlines function accepts lines and strips extraneous HTML and other tags and returns the
    cleaned text.
    """
    result_lines = []
    current_line = ""
    for line in lines:
        # line = line.strip().decode("utf-8")
        line = line.strip()
        # Convert any html elements to markdown
        line = html2markdown.convert(line)
        # Remove the vtt header
        if re.match("WEBVTT", line):
            continue
        # Skip lines that begin with a number
        if re.match(r"^\d", line):
            continue
        # Skip lines that contain only numbers and symbols
        if re.match(r"^[0-9\W]+$", line):
            continue
        # current_line += line
        # Strip any remaining html stuff
        line = re.sub("<[^<]+?>", "", line)
        # Skip blank lines
        if not line:
            continue
        # Concatenate lines less than 120 characters
        # NOTE: This is the lazy way, really should be working with words or sentences.
        if line:
            if len(current_line) + len(line) <= 120:
                current_line += line + " "
            else:
                result_lines.append(current_line.strip())
                current_line = line + " "
    # Add the last line
    if current_line:
        result_lines.append(current_line.strip())
    return (result_lines)


def savedocx(title, subs):
    """
    The savedocx funtion accepts two parameters: title and subtitles.  It creates a new docx file named
    [title].docx with [title] as a heading and [subs] as one contiguous paragraph in the document.
    """
    savename = title + ".docx"
    print("Saving subtitles in docx format to: ")
    print(savename)
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("\n".join(subs))
    document.save(savename)
    return ()


def savetext(title, subs):
    """
    The savetext function accepts two parameters: title and subtitles.  It creates a new text file named
    [title].txt containing [subs].
    """
    print("Saving subtitles in text to: ")
    savename = title + ".txt"
    print(savename)
    f = open(savename, "w", encoding="utf8")
    # f.writelines(subs)
    f.write("\n".join(subs))
    f.close()
    return ()

def webmenu():
    """
    The webmenu function does not accept any parameters. It queries the currently available videos from {url},
    and presents a menu for the user to select from. Upon selection, the video subtitles are displayed.
    """
    r = requests.get(url)
    r_dict = r.json()
    r_json = json.dumps(r_dict, indent=4)
    parsed_json = json.loads(r_json)
    media = parsed_json["category"]["media"]
    vidslist = []
    ids = 0
    for item in range(len(media)):
        print("Item:", item)
        video_key = media[item]["languageAgnosticNaturalKey"]
        title = media[item]["title"]
        if "subtitles" not in media[item]["files"][3]:
            continue
        subtitles_url = media[item]["files"][3]["subtitles"]["url"]
        video = media[item]["files"][3]["progressiveDownloadURL"]
        download = download_file(subtitles_url)
        subs = processlines(download)
        keys = ['id', 'video key', 'name', 'link', 'subtitles']
        values = [ids, video_key, title, video, subs]
        vidslist.append(values)
        ids = ids + 1
    while True:
        print("\n", banner)
        print("Number \t Title")
        print("------ \t -----")
        for item in range(len(vidslist)):
            print(vidslist[item][0], "\t", vidslist[item][2])
        print("\n")
        vidnum = input(
            "Please enter the video number, supply the URL of the video file (must end with .mp4) or q to exit: ")
        if vidnum == "":
            continue
        if vidnum == "q":
            break
        if vidnum.startswith('http') and vidnum.endswith('.mp4'):
            download = download_video(vidnum, tmp)
            if download == "fail":
                print("I'm sorry, the download failed.")
                continue
            title = get_title(download)
            subs = get_subtitles(download)
            subs = processlines(subs)
            print("\nTitle: ", title)
            print("\nNow finding paragraphs.....")
            paragraphs = break_into_paragraphs_nltk(str(subs))
            for i, paragraph in enumerate(paragraphs):
                print(f"  {paragraph}\n  ")
        else:
            if not vidnum.strip().isdigit():
                print("ERROR: Please enter a valid number or URL.")
                continue
            if int(vidnum) > len(vidslist) - 1:
                print("Error!\n  Please enter a number from the available options.")
                print("  Maximum number is:", len(vidslist) - 1)
                continue
            print("Processing entry ", vidnum)
            subtitle = vidslist[int(vidnum)]
            print("\nTitle: ", subtitle[2])
            print("Link: ", subtitle[3], "\n")
            paragraphs = break_into_paragraphs_nltk(str(subtitle[4]))  # Use subtitle[4] instead of subs
            for i, paragraph in enumerate(paragraphs):
                print(f"  {paragraph}\n  ")
    return ()
# def webmenu():
#     """
#     The webmenu function does not accept any parameters.  It queries the currently available videos from {url},
#     and presents a menu for the user to select from.  Upon selection, the video subtitles are displayed."""
#     r = requests.get(url)
#     r_dict = r.json()
#     r_json = json.dumps(r_dict, indent=4)
#     parsed_json = json.loads(r_json)
#     media = parsed_json["category"]["media"]
#     vidslist = []
#     ids = 0
#     for item in range(len(media)):
#         print("Item:", item)
#         video_key = media[item]["languageAgnosticNaturalKey"]
#         # print(video_key)
#         title = media[item]["title"]
#         # Added logic for video files that do not contain subtitles like song releases
#         if "subtitles" not in media[item]["files"][3]:
#             continue
#         subtitles_url = media[item]["files"][3]["subtitles"]["url"]
#         # print(subtitles_url)
#         video = media[item]["files"][3]["progressiveDownloadURL"]
#         download = download_file(subtitles_url)
#         subs = processlines(download)
#         keys = ['id', 'video key', 'name', 'link', 'subtitles']
#         values = [ids, video_key, title, video, subs]
#         vidslist.append(values)
#         ids = ids + 1
#         # print(video_key)
#     # print("Number of videos available: ", len(vidslist))
#     # print(vidslist)
#     # exit()
#     while True:
#         print("\n", banner)
#         print("Number \t Title")
#         print("------ \t -----")
#         for item in range(len(vidslist)):
#             print(vidslist[item][0], "\t", vidslist[item][2])
#             # print(vidslist[item][0] + 1, "\t", vidslist[item][2])
#         print("\n")
#         vidnum = input(
#             "Please enter the video number, supply the URL of the video file (must end with .mp4) or q to exit: ")
#         if vidnum == "":
#             continue
#         if vidnum == "q":
#             break
#         if vidnum.startswith('http') and vidnum.endswith('.mp4'):
#             download = download_video(vidnum, tmp)
#             if download == "fail":
#                 print("I'm sorry, the download failed.")
#                 continue
#             title = get_title(download)
#             subs = get_subtitles(download)
#             subs = processlines(subs)
#             print("\nTitle: ", title)
#             # print("\n-----Subtitles begin here.-----\n")
#             # print(*subs, sep="\n")
#             print("\nNow finding paragraphs.....")
#             print(break_into_paragraphs_nltk(subs))
#             # print("\n-----Subtitles end here.-----\n")
#         # else:
#         #    print("ERROR: That doesn't appear to be a valid URL.\n")
#         if not vidnum.strip().isdigit():
#             continue
#         if int(vidnum) > len(vidslist) - 1:
#             print("Error!\n  Please enter a number from the available options.")
#             print("  Maximum number is:", len(vidslist) - 1)
#             continue
#         print("Processing entry ", vidnum)
#         subtitle = vidslist[int(vidnum)]
#         print("\nTitle: ", subtitle[2])
#         print("Link: ", subtitle[3], "\n")
#         # print("-----Subtitles begin here.-----\n")
#         # print(*subtitle[4], sep='\n')
#         # print(break_into_paragraphs_nltk(str(subtitle[4])))
#         paragraphs = break_into_paragraphs_nltk(str(subs))
#         # paragraphs = break_into_paragraphs_nltk(str(subtitle[4]))
#         for i, paragraph in enumerate(paragraphs):
#             print(f"  {paragraph}\n  ")
#         # print("\n-----Subtitles end here.-----\n")
#     return ()


def main():
    """
This program has two modes of operation:
    Mode 1: No command line parameters
    The program queries for the latest videos from jw.org and allows
    you to choose which one you'd like the subtitles for and then displays them.

    Mode 2:
    If you provide the URL of a video file as a command line parameter, the program downloads the video
    and attempts to extract the subtitle information from the video to display.

    There are two additional parameters that can passed when providing the URL:
        -t Outputs the subtitles to a text file in the current directory named with the
            title of the video.
        -d Outputs the subtitles to a Microsoft Word docx file in the current directory
            named with the title of the video.

    IMPORTANT: The second option requires that you have installed FFpmeg and the ffmpeg program is
    available in your PATH.

    You can find FFmpeg at:  https://ffmpeg.org/download.html
    """
    print("\nSubtitle downloader for jw.org videos.\n")
    # print(banner)
    print("Visit the official website of Jehovah's Witnesses: https://jw.org")
    print("Finding current videos.....")
    parser = argparse.ArgumentParser(
        description=main.__doc__, formatter_class=RawTextHelpFormatter)
    parser.add_argument(
        '-d',
        '--docx',
        required=False,
        action="store_true",
        help="Save subtitles to a docx file."
    )
    parser.add_argument(
        '-t',
        '--textfile',
        required=False,
        action="store_true",
        help="Save subtitles to a text file."
    )
    parser.add_argument(
        'video',
        nargs='?',
        type=str,
        default=False)
    args = parser.parse_args()
    if args.video:
        vidurl = args.video
        if vidurl.endswith('.mp4'):
            if vidurl.startswith('http'):
                download = download_video(vidurl, tmp)
            else:
                download = vidurl
            if download == "fail":
                print("I'm sorry, the download failed.")
                exit(1)
            title = get_title(download)
            subs = get_subtitles(download)
            subs = processlines(subs)
            # os.remove(srt_file)
            # Change to save to text or docx
            if args.textfile:
                savetext(title, subs)
            if args.docx:
                savedocx(title, subs)
            if not args.textfile and not args.docx:
                print("Subtitles from supplied URL.\n")
                print("\nTitle: ", title)
                print("\n-----Subtitles begin here.-----\n")
                paragraphs = break_into_paragraphs_nltk(str(subs))
                for i, paragraph in enumerate(paragraphs):
                    print(f"  {paragraph}\n  ")
                # print("Raw text")
                # print(*subs, sep="\n")
                # print("now attempting to print with paragraphs.")

                # , sep="\n")
                print("\n\n-----Subtitles end here.-----\n")
        else:
            print("ERROR: That doesn't appear to be a valid URL.\n")
            help()
    else:
        webmenu()
    print("Hope that helps.\nGoodbye.")


if __name__ == "__main__":
    main()
